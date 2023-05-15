import os
import pdf2image
import docx
import sys
from docx.shared import Mm
from docx.enum.section import WD_ORIENT
from tqdm import tqdm
import argparse


# 创建解析器对象
parser = argparse.ArgumentParser(description='Convert an pdf file to an docx file')

# 添加命令行参数
parser.add_argument('input', help='Input file name')
parser.add_argument('--output', default="a.docx", help='Output file name (defaut is a.docx)')
parser.add_argument('--dpi', type=int, default=400, help='DPI for the output file (default: 400)')


# 解析命令行参数
args = parser.parse_args()

pdf_path = args.input
docx_path = args.output
dpi = args.dpi # 指定输出图像的dpi值

# 保存中间图片位置
output_folder = 'temp_output'
if not os.path.exists(output_folder):
    os.mkdir(output_folder)

# 指定要转换为的图片格式，如'jpeg'、'png'等
image_format = 'jpeg'

# 指定Word文档的页面大小（A4纸）
page_height = 297
page_width = 210

# 指定Word文档的页边距（上下左右均为0）
margin_top = 0
margin_bottom = 0
margin_left = 0
margin_right = 0

# 指定每个图片占用一个页面
para_before = 0
para_after = page_height


# 先创建一个空白的Word文档，并设置页面大小和页边距
doc = docx.Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.PORTRAIT  # 垂直方向
section.page_height = Mm(page_height)
section.page_width = Mm(page_width)
section.top_margin = Mm(margin_top)
section.bottom_margin = Mm(margin_bottom)
section.left_margin = Mm(margin_left)
section.right_margin = Mm(margin_right)

# 使用pdf2image库将pdf文件转换为图像
try:
    print("Converting PDF to Images...", end=" ")
    images = pdf2image.convert_from_path(
        pdf_path,
        dpi=dpi,  # 设置dpi参数
        grayscale=False,
        thread_count=5)
    print("done")
except:
    print("\n")
    print("convert error")
    sys.exit(1)
    
    
    

# 并遍历每一页图像并且将其插入到word中
for i, image in enumerate(tqdm(images, desc='Converting Images to Docx')):
    # 将图像保存为指定格式的文件
    image_filename = os.path.join(output_folder, f'page{i + 1}.{image_format}')
    image.save(image_filename, image_format)

    # 将图像插入到Word文档中的单独段落，并设置段前/后距离以使其占用一个页面
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Mm(para_before)
    paragraph.paragraph_format.space_after = Mm(para_after)
    run = paragraph.add_run()
    run.add_picture(image_filename, width=Mm(page_width))  # 指定图片宽度为页面宽度

# 保存修改后的Word文档
doc.save(docx_path)
